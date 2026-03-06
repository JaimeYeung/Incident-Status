"""Convert PRD.md to PRD.docx."""
from docx import Document
from docx.shared import Pt

doc = Document()
with open("PRD.md", "r", encoding="utf-8") as f:
    content = f.read()

for line in content.split("\n"):
    s = line.strip()
    if not s:
        continue
    if s.startswith("# "):
        doc.add_heading(s[2:], level=0)
    elif s.startswith("## "):
        doc.add_heading(s[3:], level=1)
    elif s.startswith("### "):
        doc.add_heading(s[4:], level=2)
    elif s.startswith("- "):
        doc.add_paragraph(s[2:], style="List Bullet")
    elif s.startswith("|"):
        doc.add_paragraph(s)
    else:
        doc.add_paragraph(s)

doc.save("PRD.docx")
print("Saved PRD.docx")
